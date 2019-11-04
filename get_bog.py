import os
from time import sleep
import unicodedata
import pyperclip
import pyautogui, sys
from docx import Document
import json


def iterate_banks():
    docx_path = 'C:\\Users\\john\\Desktop\\banks_working_data\\'
    pyautogui.FAILSAFE = False
    
    jsons_dir = 'bank_jsons'
    json_files = os.listdir(jsons_dir)
    for json_file in json_files:
        if json_file == '.DS_Store':
            pass
        else:
            file_path = os.path.join(jsons_dir, json_file)
            print(file_path)
            with open(file_path, 'r') as fin:
                data = json.load(fin)

                bank_name = data['name']
                bank_all_text = data['master_string']
                bank_refined_text = data['refined_master_string']
                #bank_num_pages = data['num_pages']

                if len(bank_refined_text) < 10:
                    with open('already_processed.txt', 'a') as fout:
                        fout.write(json_file + '\n')

                else:

                    # process 'all_text'
                    bank_file_name = bank_name.replace(' ', '_') + '_all' + '.docx'
                    document = Document()
                    document.add_paragraph(bank_all_text)
                    document.save(bank_file_name)

                    document_length = bank_all_text.split()
                    sleep_timer = (len(document_length) / 1000) * 10
                    print(sleep_timer)

                    # process chunked text if too large for stylewriter 
                    if sleep_timer > 6500:
                        os.remove(docx_path + bank_file_name)
                        text_chunks = [document_length[x:x+350000] for x in range(0, len(document_length), 350000)]
                        for i in range(len(text_chunks)):
                            print('Working on chunk ' + str(i))
                            chunk_text = ' '.join(text_chunks[i])
                            sleep_timer = (len(text_chunks[i]) / 1000) * 10
                            chunk_document_name = str(i) + '_' + bank_file_name
                            document = Document()
                            document.add_paragraph(chunk_text)
                            document.save(chunk_document_name)
                                
                            os.startfile(docx_path + chunk_document_name)
                            sleep(1.2)
                            pyautogui.click(849, 63)
                            sleep(1)
                            pyautogui.press(['down', 'enter'])
                            sleep(sleep_timer)
                            sleep(15)
                            pyautogui.click(884, 65)
                            sleep(1)
                            pyautogui.click(811, 16)
                            sleep(.5)
                            pyautogui.press(['right', 'enter'])
                            sleep(1)
                            os.remove(docx_path + chunk_document_name)
                            sleep(1)

                        sleep(1)
                        print()
                    
                    # process single text
                    else:
                        os.startfile(docx_path + bank_file_name)
                        sleep(1.2)
                        pyautogui.click(849, 63)
                        sleep(1)
                        pyautogui.press(['down', 'enter'])
                        sleep(sleep_timer)
                        sleep(15)
                        pyautogui.click(884, 65)
                        sleep(1)
                        pyautogui.click(811, 16)
                        sleep(.5)
                        pyautogui.press(['right', 'enter'])
                        sleep(1)
                        os.remove(docx_path + bank_file_name)
                        sleep(1)

                    # process 'refined_text'
                    bank_file_name = bank_name.replace(' ', '_') + '_refined' + '.docx'
                    document = Document()
                    document.add_paragraph(bank_refined_text)
                    document.save(bank_file_name)

                    document_length = bank_refined_text.split()
                    sleep_timer = (len(document_length) / 1000) * 10
                    print(sleep_timer)

                    # process chunked text if too large for stylewriter 
                    if sleep_timer > 6500:
                        os.remove(docx_path + bank_file_name)
                        text_chunks = [document_length[x:x+350000] for x in range(0, len(document_length), 350000)]
                        for i in range(len(text_chunks)):
                            print('Working on chunk ' + str(i))
                            chunk_text = ' '.join(text_chunks[i])
                            sleep_timer = (len(text_chunks[i]) / 1000) * 10
                            chunk_document_name = str(i) + '_' + bank_file_name
                            document = Document()
                            document.add_paragraph(chunk_text)
                            document.save(chunk_document_name)
                                
                            os.startfile(docx_path + chunk_document_name)
                            sleep(1.2)
                            pyautogui.click(849, 63)
                            sleep(1)
                            pyautogui.press(['down', 'enter'])
                            sleep(sleep_timer)
                            sleep(15)
                            pyautogui.click(884, 65)
                            sleep(1)
                            pyautogui.click(811, 16)
                            sleep(.5)
                            pyautogui.press(['right', 'enter'])
                            sleep(1)
                            os.remove(docx_path + chunk_document_name)
                            sleep(1)

                        sleep(1)
                        print()
                    
                    # process single text
                    else:
                        os.startfile(docx_path + bank_file_name)
                        sleep(1.2)
                        pyautogui.click(849, 63)
                        sleep(1)
                        pyautogui.press(['down', 'enter'])
                        sleep(sleep_timer)
                        sleep(15)
                        pyautogui.click(884, 65)
                        sleep(1)
                        pyautogui.click(811, 16)
                        sleep(.5)
                        pyautogui.press(['right', 'enter'])
                        sleep(1)
                        os.remove(docx_path + bank_file_name)
                        sleep(1)

                    with open('already_processed.txt', 'a') as fout:
                        fout.write(json_file + '\n')


def get_spot():
    try:
        while True:
            x, y = pyautogui.position()
            positionStr = 'X: ' + str(x).rjust(4) + ' Y: ' + str(y).rjust(4)
            print(positionStr, end='')
            print('\b' * len(positionStr), end='', flush=True)
    except KeyboardInterrupt:
        print('\n')


iterate_banks()
#get_spot()

